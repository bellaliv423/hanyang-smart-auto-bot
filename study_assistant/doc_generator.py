"""
Word 문서 자동 생성기 - 수업 주제 → 전문적인 .docx 파일

기능:
1. 주제 입력 → Claude AI가 구조화된 보고서/레포트 생성
2. 참고 논문/기사/사례 포함
3. 한국어/中文/English 다국어 지원
4. Google Drive 자동 업로드

사용법:
    python study_assistant/doc_generator.py --topic "회귀분석의 기초" --course 196594
    python study_assistant/doc_generator.py --topic "M&A 사례분석" --course 196656 --type report
    python study_assistant/doc_generator.py --topic "계약법 요약" --lang zh --type summary
"""

import anthropic
import argparse
import io
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

PROJECT_ROOT = Path(__file__).parent.parent

from dotenv import load_dotenv
load_dotenv(PROJECT_ROOT / "config" / ".env")

API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

COURSE_MAP = {
    196594: {"ko": "경영통계학", "en": "Business Statistics", "zh": "經營統計學", "prof": "부제만"},
    196600: {"ko": "상법및계약법", "en": "Commercial & Contract Law", "zh": "商法及契約法", "prof": "강편모"},
    196656: {"ko": "M&A전략", "en": "M&A Strategy", "zh": "M&A戰略", "prof": "김철중"},
    196622: {"ko": "국제거시금융론", "en": "International Macrofinance", "zh": "國際宏觀金融論", "prof": "이창민"},
}

DOC_TYPES = {
    "report": "학술 보고서/레포트",
    "summary": "수업 내용 요약",
    "case": "기업 사례 분석",
    "review": "논문/자료 리뷰",
    "essay": "에세이/소논문",
}


def generate_doc_content(topic, course_id=None, lang="ko", doc_type="report"):
    """Claude AI로 문서 내용 생성"""
    client = anthropic.Anthropic(api_key=API_KEY, timeout=120.0)

    course_info = ""
    if course_id and course_id in COURSE_MAP:
        c = COURSE_MAP[course_id]
        course_info = f"\n과목: {c['ko']} ({c['en']}) / 교수: {c['prof']}"

    lang_map = {"ko": "한국어", "en": "English", "zh": "中文繁體"}
    target_lang = lang_map.get(lang, "한국어")
    type_desc = DOC_TYPES.get(doc_type, "보고서")

    prompt = f"""MBA 수업용 {type_desc}를 작성해주세요.

주제: {topic}{course_info}
언어: {target_lang}
문서 유형: {type_desc}

반드시 아래 JSON 형식으로 출력하세요:

```json
{{
  "title": "문서 제목",
  "subtitle": "부제 (선택)",
  "abstract": "요약 (2-3문장)",
  "sections": [
    {{
      "heading": "1. 서론",
      "content": "본문 내용 (상세하게 작성)",
      "subsections": [
        {{
          "heading": "1.1 배경",
          "content": "하위 섹션 내용"
        }}
      ]
    }},
    {{
      "heading": "2. 본론",
      "content": "본문 내용"
    }},
    {{
      "heading": "3. 사례 연구",
      "content": "기업 사례 분석"
    }},
    {{
      "heading": "4. 결론",
      "content": "결론 및 시사점"
    }}
  ],
  "references": [
    "저자명 (연도). 제목. 출판사/저널.",
    "기사/논문 출처"
  ],
  "key_terms": [
    {{"term": "용어", "ko": "한국어", "en": "English", "zh": "中文"}}
  ]
}}
```

요구사항:
1. MBA 수준의 전문적이고 학술적인 내용
2. 각 섹션은 300자 이상 상세하게 작성
3. 실제 기업 사례 포함
4. 관련 논문/기사 참고자료 최소 3개
5. 핵심 용어는 한/영/중 3개 언어 병기"""

    response = client.messages.create(
        model="claude-sonnet-4-5-20250929",
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}],
    )

    text = response.content[0].text

    json_match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if json_match:
        return json.loads(json_match.group(1))

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return {
            "title": topic,
            "subtitle": "",
            "abstract": "",
            "sections": [{"heading": "내용", "content": text[:3000]}],
            "references": [],
            "key_terms": []
        }


def setup_styles(doc):
    """문서 스타일 설정"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)

    for level in range(1, 4):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = 'Malgun Gothic'
            h_style.font.color.rgb = RGBColor(0, 51, 102)


def build_doc(data, course_id=None, doc_type="report", output_path=None):
    """Word 문서 생성"""
    doc = Document()
    setup_styles(doc)

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 제목
    title_para = doc.add_heading(data["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 부제
    if data.get("subtitle"):
        sub = doc.add_paragraph(data["subtitle"])
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(14)
        sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # 과목 + 날짜
    if course_id and course_id in COURSE_MAP:
        c = COURSE_MAP[course_id]
        info_text = f"{c['ko']} | {c['prof']} 교수 | {datetime.now().strftime('%Y년 %m월 %d일')}"
    else:
        info_text = f"한양대학교 MBA | {datetime.now().strftime('%Y년 %m월 %d일')}"

    info = doc.add_paragraph(info_text)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(10)
    info.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph("")  # 공백

    # 요약
    if data.get("abstract"):
        doc.add_heading("Abstract / 요약", level=2)
        abstract_para = doc.add_paragraph(data["abstract"])
        abstract_para.runs[0].font.italic = True
        doc.add_paragraph("")

    # 본문 섹션
    for section in data.get("sections", []):
        doc.add_heading(section["heading"], level=1)
        if section.get("content"):
            doc.add_paragraph(section["content"])

        for subsection in section.get("subsections", []):
            doc.add_heading(subsection["heading"], level=2)
            if subsection.get("content"):
                doc.add_paragraph(subsection["content"])

    # 핵심 용어
    key_terms = data.get("key_terms", [])
    if key_terms:
        doc.add_heading("Key Terms / 핵심 용어 / 關鍵術語", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = "한국어"
        hdr[1].text = "English"
        hdr[2].text = "中文"
        for term in key_terms:
            row = table.add_row().cells
            row[0].text = term.get("ko", "")
            row[1].text = term.get("en", "")
            row[2].text = term.get("zh", "")

    # 참고자료
    refs = data.get("references", [])
    if refs:
        doc.add_heading("References / 참고 자료", level=1)
        for i, ref in enumerate(refs, 1):
            doc.add_paragraph(f"[{i}] {ref}")

    # 푸터
    doc.add_paragraph("")
    footer = doc.add_paragraph("Generated by Hanyang Smart Auto Study Bot")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(180, 180, 180)

    # 저장
    if not output_path:
        safe_title = data["title"].replace(" ", "_").replace("/", "-")[:30]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        output_path = PROJECT_ROOT / "data" / "study_notes" / f"DOC_{safe_title}_{timestamp}.docx"

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return output_path


def main():
    parser = argparse.ArgumentParser(description="Smart Auto Document Generator")
    parser.add_argument("--topic", required=True, help="문서 주제")
    parser.add_argument("--course", type=int, help="Course ID")
    parser.add_argument("--lang", default="ko", choices=["ko", "en", "zh"], help="언어")
    parser.add_argument("--type", default="report", choices=DOC_TYPES.keys(), help="문서 유형")
    parser.add_argument("--output", help="출력 파일 경로")
    parser.add_argument("--upload", action="store_true", help="Drive 업로드")
    parser.add_argument("--email", action="store_true", help="이메일로 발송")
    parser.add_argument("--whatsapp", action="store_true", help="WhatsApp 알림")
    parser.add_argument("--send", action="store_true", help="이메일 + WhatsApp 모두 발송")
    args = parser.parse_args()

    if not API_KEY:
        print("[!] ANTHROPIC_API_KEY 필요!")
        sys.exit(1)

    print(f"{'='*50}")
    print(f"  Smart Auto Document Generator")
    print(f"  주제: {args.topic}")
    print(f"  유형: {DOC_TYPES[args.type]}")
    print(f"{'='*50}")

    # 1. AI 내용 생성
    print("\n[1/3] AI 내용 생성 중...")
    data = generate_doc_content(args.topic, args.course, args.lang, args.type)
    print(f"  제목: {data.get('title', args.topic)}")
    print(f"  섹션: {len(data.get('sections', []))}개")

    # 2. 문서 생성
    print("\n[2/3] Word 문서 생성 중...")
    output_path = build_doc(data, args.course, args.type, args.output)
    print(f"  저장: {output_path}")
    print(f"  크기: {output_path.stat().st_size / 1024:.1f} KB")

    # 3. Drive 업로드
    if args.upload:
        print("\n[3/3] Google Drive 업로드 중...")
        try:
            sys.path.insert(0, str(PROJECT_ROOT / "scrapers"))
            from drive_uploader import get_drive_service, find_or_create_folder, upload_file

            service = get_drive_service()
            DRIVE_ROOT = "1RGGiiz_DKf5ZcUNk7baJis95oAcDJ5Em"

            with open(PROJECT_ROOT / "config" / "courses.json", "r", encoding="utf-8") as f:
                config = json.load(f)

            semester_folder = find_or_create_folder(service, f"한양MBA_{config['semester']}", DRIVE_ROOT)

            if args.course and args.course in COURSE_MAP:
                course_name = COURSE_MAP[args.course]["ko"]
                course_folder = find_or_create_folder(service, course_name, semester_folder)
                notes_folder = find_or_create_folder(service, "학습노트", course_folder)
                file_id = upload_file(service, output_path, notes_folder)
            else:
                file_id = upload_file(service, output_path, semester_folder)

            print(f"  Drive: https://drive.google.com/file/d/{file_id}/view")
        except Exception as e:
            print(f"  [!] 업로드 실패: {e}")

    # 4. 이메일 발송
    if args.email or args.send:
        print("\n[EMAIL] 이메일 발송 중...")
        try:
            from email_sender import send_email_with_attachment
            course_name = COURSE_MAP.get(args.course, {}).get("ko", "") if args.course else ""
            subject = f"[Smart Auto] {course_name} - {data.get('title', args.topic)}"
            send_email_with_attachment(output_path, subject=subject)
        except Exception as e:
            print(f"  [!] 이메일 실패: {e}")

    # 5. WhatsApp 알림
    if args.whatsapp or args.send:
        print("\n[WHATSAPP] WhatsApp 알림 발송 중...")
        try:
            import subprocess
            target = os.getenv("WHATSAPP_TARGET", "")
            course_name = COURSE_MAP.get(args.course, {}).get("ko", "") if args.course else ""
            msg = (
                f"문서 생성 완료!\n"
                f"제목: {data.get('title', args.topic)}\n"
                f"과목: {course_name}\n"
                f"유형: {DOC_TYPES[args.type]}\n"
                f"파일: {output_path.name}\n"
                f"이메일로도 보냈어요!"
            )
            subprocess.run(
                ["wsl", "bash", "-c",
                 f'npx openclaw message send --channel whatsapp '
                 f'--target "{target}" --message "{msg}" --json'],
                capture_output=True, timeout=30,
            )
            print("  WhatsApp 알림 발송 완료!")
        except Exception as e:
            print(f"  [!] WhatsApp 실패: {e}")

    print(f"\n{'='*50}")
    print(f"[DONE] 문서 생성 완료!")

    result = {
        "status": "success",
        "file": str(output_path),
        "file_name": output_path.name,
        "title": data.get("title", args.topic),
        "sections": len(data.get("sections", [])),
    }
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
